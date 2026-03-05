"""
Flask Backend API for Phishing URL Detection System
Provides endpoints for URL scanning, QR-based scanning, DOCX reports,
reporting to cybercrime, and statistics.
"""

import os
import re
import io
import json
import time
from datetime import datetime
from urllib.parse import urlparse
import csv

from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from dotenv import load_dotenv
import joblib
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from feature_extractor import extract_features, features_to_vector, FEATURE_NAMES
from sqlite_db import (
    init_db, save_scan, save_report,
    get_recent_scans, get_scan_stats, get_reports, get_all_scans
)

load_dotenv()

# ──────────────────────────────────────────────────────────────
# App Configuration
# ──────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder='../frontend', static_url_path='')
CORS(app, resources={r"/api/*": {"origins": "*"}})

limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per hour", "50 per minute"],
    storage_uri="memory://"
)

# ──────────────────────────────────────────────────────────────
# Load ML Model
# ──────────────────────────────────────────────────────────────
MODEL_PATH = os.path.join(os.path.dirname(__file__), 'model', 'phishing_model.joblib')
model = None

LABEL_MAP = {0: 'Safe', 1: 'Suspicious', 2: 'Phishing'}

# In-memory storage (fallback when Firebase is offline)
in_memory_scans = []
in_memory_reports = []
in_memory_stats = {'total_scans': 0, 'phishing_count': 0, 'suspicious_count': 0, 'safe_count': 0}


def load_model():
    """Load the trained ML model."""
    global model
    if os.path.exists(MODEL_PATH):
        model = joblib.load(MODEL_PATH)
        print(f"[INFO] ML model loaded from {MODEL_PATH}")
    else:
        print(f"[WARNING] Model not found at {MODEL_PATH}")
        print("  Run 'python train_model.py' first to train the model.")


# ──────────────────────────────────────────────────────────────
# Input Validation & Sanitization
# ──────────────────────────────────────────────────────────────
def sanitize_input(text: str) -> str:
    """Sanitize input to prevent XSS attacks."""
    if not text:
        return ''
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&#x27;')
    return text.strip()


def is_valid_url(url: str) -> bool:
    """Validate URL format."""
    if not url or len(url) > 2048:
        return False
    if not url.startswith(('http://', 'https://', 'ftp://')):
        url = 'http://' + url
    try:
        result = urlparse(url)
        return bool(result.netloc)
    except Exception:
        return False


# ──────────────────────────────────────────────────────────────
# Core Scan Logic (shared by URL scan and QR scan)
# ──────────────────────────────────────────────────────────────
def perform_scan(url: str, source: str = 'URL Input') -> dict:
    """Run feature extraction and ML prediction on a URL."""
    features = extract_features(url)
    feature_vector = features_to_vector(features)

    if model is None:
        raise RuntimeError('ML model not loaded. Please train the model first.')

    feature_array = np.array([feature_vector])
    prediction_id = model.predict(feature_array)[0]
    probabilities = model.predict_proba(feature_array)[0]

    prediction = LABEL_MAP.get(int(prediction_id), 'Unknown')
    confidence = round(float(max(probabilities)) * 100, 2)

    risk_score = round(
        (probabilities[1] * 50 + probabilities[2] * 100), 2
    )
    risk_score = min(max(risk_score, 0), 100)

    timestamp = datetime.utcnow().isoformat() + 'Z'

    result = {
        'url': url,
        'prediction': prediction,
        'confidence': confidence,
        'risk_score': risk_score,
        'features': features,
        'timestamp': timestamp,
        'scan_source': source,
        'probabilities': {
            'safe': round(float(probabilities[0]) * 100, 2),
            'suspicious': round(float(probabilities[1]) * 100, 2),
            'phishing': round(float(probabilities[2]) * 100, 2)
        }
    }

    # Save to Firebase (or in-memory)
    doc_id = save_scan(result)
    if doc_id:
        result['scan_id'] = doc_id
    else:
        in_memory_scans.insert(0, result)
        if len(in_memory_scans) > 100:
            in_memory_scans.pop()
        in_memory_stats['total_scans'] += 1
        if prediction == 'Phishing':
            in_memory_stats['phishing_count'] += 1
        elif prediction == 'Suspicious':
            in_memory_stats['suspicious_count'] += 1
        else:
            in_memory_stats['safe_count'] += 1

    return result


# ──────────────────────────────────────────────────────────────
# Email Analysis Logic
# ──────────────────────────────────────────────────────────────
PHISHING_KEYWORDS = [
    'verify', 'bank', 'account', 'suspend', 'security', 'password', 'login', 'billing',
    'payment', 'urgent', 'immediately', 'soon', 'action required', 'alert', 'important',
    'update', 'confirm', 'customer support', 'help desk', 'officer', 'department',
    'tax', 'irs', 'police', 'lottery', 'winner', 'gift card', 'inheritance', 'amazon',
    'netflix', 'paypal', 'microsoft', 'apple', 'office 365', 'itunes', 'wallet', 'crypto'
]

URGENCY_PHRASES = [
    '24 hours', '48 hours', 'limit exceeded', 'unauthorized access',
    'blocked', 'restricted', 'legal action', 'final warning', 'locked',
    'unusual activity', 'signed in on new device', 'avoid suspension'
]

def analyze_email_text(text: str) -> dict:
    """Extract links and analyze text for phishing indicators."""
    # Find URLs
    url_pattern = r'https?://[^\s<>"]+|www\.[^\s<>"]+'
    found_urls = list(set(re.findall(url_pattern, text))) # Unique URLs
    
    # Analyze text keywords
    text_lower = text.lower()
    flags = []
    text_score = 0
    
    keyword_matches = [k for k in PHISHING_KEYWORDS if k in text_lower]
    if keyword_matches:
        text_score += min(len(keyword_matches) * 8, 40)
        flags.append(f"Suspicious Brand/Service Mentioned: {', '.join(keyword_matches[:3])}")
        
    urgency_matches = [p for p in URGENCY_PHRASES if p in text_lower]
    if urgency_matches:
        text_score += 20
        flags.append(f"High-Urgency Threat Pattern: '{urgency_matches[0]}'")
        
    # Potential spoofing indicators
    if 'dear customer' in text_lower or 'dear user' in text_lower:
        text_score += 15
        flags.append("Generic Non-Personalized Greeting")
        
    if 'click here' in text_lower or 'click the link' in text_lower:
        text_score += 10
        flags.append("Direct Call-to-Action for External Link")

    # Scan URLs found
    url_results = []
    max_url_score = 0
    for url in found_urls[:8]: # Scan more URLs
        if not url.startswith('http'):
            url = 'http://' + url
        try:
            res = perform_scan(url, source='Email Content')
            url_results.append({
                'url': url,
                'prediction': res['prediction'],
                'risk_score': res['risk_score']
            })
            max_url_score = max(max_url_score, res['risk_score'])
        except:
            continue

    if max_url_score > 70:
        flags.append(f"CRITICAL: High-Risk Malicious Link Detected")
    elif max_url_score > 30:
        flags.append(f"Suspicious Link Detected in Content")

    # Combine scores
    final_score = min(max_url_score + (text_score * 0.6), 100)
    
    if final_score > 75 or max_url_score > 85:
        prediction = 'Phishing'
    elif final_score > 40:
        prediction = 'Suspicious'
    else:
        prediction = 'Safe'
        
    # Detect HTTPS Trust Trap (adapted for email analysis context)
    # This logic is typically applied during feature extraction for a single URL.
    # For email analysis, we'll check if any of the found URLs exhibit this pattern.
    is_ssl_trap_overall = False
    for url_res in url_results:
        url_obj = urlparse(url_res['url'])
        if url_obj.scheme == 'https':
            # High-risk TLDs commonly used with free SSL for phishing
            suspicious_tld_list = ['.xyz', '.top', '.ga', '.cf', '.ml', '.tk', '.date', '.icu']
            if any(url_obj.netloc.lower().endswith(tld) for tld in suspicious_tld_list):
                is_ssl_trap_overall = True
                flags.append(f"HTTPS Trust Trap: Suspicious TLD '{url_obj.netloc.split('.')[-1]}' with HTTPS for {url_res['url']}")
                break # Flag once if found

            # Brand name in domain but not as the main domain
            brands = ['amazon', 'google', 'microsoft', 'netflix', 'paypal', 'bank', 'apple', 'office']
            domain_parts = url_obj.netloc.lower().split('.')
            if len(domain_parts) > 2:
                main_domain = domain_parts[-2]
                for brand in brands:
                    if brand in url_obj.netloc.lower() and brand != main_domain:
                        is_ssl_trap_overall = True
                        flags.append(f"HTTPS Trust Trap: Brand Impersonation '{brand}' in subdomain/path with HTTPS for {url_res['url']}")
                        break
        if is_ssl_trap_overall:
            break

    if is_ssl_trap_overall:
        # If an SSL trap is detected, increase the final score and potentially change prediction
        final_score = min(final_score + 25, 100) # Significant penalty
        if final_score > 75:
            prediction = 'Phishing'
        elif final_score > 40 and prediction == 'Safe': # If it was safe, now it's at least suspicious
            prediction = 'Suspicious'
        
    timestamp = datetime.utcnow().isoformat() + 'Z'
    
    res_dict = {
        'url': f"Email Content ({len(found_urls)} links)",
        'prediction': prediction,
        'risk_score': round(final_score, 2),
        'confidence': 88.0,
        'flags': flags,
        'urls_found': url_results,
        'url_count': len(found_urls),
        'timestamp': timestamp,
        'is_email_scan': True
    }

    # Save to database for history
    save_scan({
        'url': res_dict['url'],
        'prediction': prediction,
        'confidence': 88.0,
        'risk_score': round(final_score, 2),
        'timestamp': timestamp,
        'scan_source': 'Email Analysis',
        'features': {},
        'probabilities': {}
    })

    return res_dict


# ──────────────────────────────────────────────────────────────
# DOCX Report Generation
# ──────────────────────────────────────────────────────────────
def generate_docx_report(scan_data: dict) -> io.BytesIO:
    """Generate a professional DOCX report from scan data (URL or Email)."""
    document = Document()
    is_email = scan_data.get('is_email_scan', False)

    # Define styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Segoe UI' if not is_email else 'Arial'
    font.size = Pt(11)

    # Header
    header_text = 'PhishGuard Detection Report' if not is_email else 'PhishGuard Email Security Assessment'
    header = document.add_heading(header_text, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Basic Info Table
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    def set_cell(row, label, value):
        table.rows[row].cells[0].text = label
        table.rows[row].cells[1].text = str(value)

    set_cell(0, 'Source Type:', 'URL Scan' if not is_email else 'Email Body Analysis')
    set_cell(1, 'Target/Context:', scan_data.get('url', 'N/A'))
    set_cell(2, 'Detection Result:', scan_data.get('prediction', 'Unknown').upper())
    set_cell(3, 'Timestamp (UTC):', scan_data.get('timestamp', 'N/A'))

    document.add_paragraph() # Spacer

    # Risk Assessment
    document.add_heading('Risk Assessment', level=1)
    risk_score = scan_data.get('risk_score', 0)
    prediction = scan_data.get('prediction', 'Unknown')
    
    p = document.add_paragraph()
    run = p.add_run(f"Result: {prediction.upper()} (Risk Score: {risk_score}/100)")
    run.bold = True
    if prediction.lower() == 'phishing':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif prediction.lower() == 'suspicious':
        run.font.color.rgb = RGBColor(255, 140, 0)
    else:
        run.font.color.rgb = RGBColor(0, 150, 0)

    if is_email:
        # Email Specific Content
        document.add_heading('Detected Indicators (Red Flags)', level=2)
        flags = scan_data.get('flags', [])
        if flags:
            for flag in flags:
                document.add_paragraph(f"🚩 {flag}", style='List Bullet')
        else:
            document.add_paragraph("No significant psychological red flags were detected in the text body.")

        document.add_heading('Embedded Link Analysis', level=2)
        urls = scan_data.get('urls_found', [])
        if urls:
            for u in urls:
                status = u.get('prediction', 'Unknown')
                score = u.get('risk_score', 0)
                p = document.add_paragraph(style='List Bullet')
                p.add_run(f"Link: {u.get('url')}\n").italic = True
                p.add_run(f"Analysis: {status} (Risk: {score}%)")
        else:
            document.add_paragraph("No HTTP/HTTPS links were detected in the provided content.")
    else:
        # URL Specific Content (Features)
        document.add_heading('URL Pattern Analysis', level=2)
        probs = scan_data.get('probabilities', {})
        document.add_paragraph(f"ML Confidence: {scan_data.get('confidence', 0)}%")
        document.add_paragraph(f"- Probability Safe: {probs.get('safe', 0)}%", style='List Bullet')
        document.add_paragraph(f"- Probability Suspicious: {probs.get('suspicious', 0)}%", style='List Bullet')
        document.add_paragraph(f"- Probability Phishing: {probs.get('phishing', 0)}%", style='List Bullet')

        # Features Table
        document.add_heading('Technical URL Features', level=2)
        features = scan_data.get('features', {})
        if features:
            f_table = document.add_table(rows=1, cols=2)
            f_table.style = 'Table Grid'
            hdr_cells = f_table.rows[0].cells
            hdr_cells[0].text = 'Feature'
            hdr_cells[1].text = 'Value'
            
            # Show top 10 most relevant features for the report
            for key, val in list(features.items())[:15]:
                row_cells = f_table.add_row().cells
                row_cells[0].text = key.replace('num_', '').replace('has_', '').replace('_', ' ').title()
                row_cells[1].text = str(val if not isinstance(val, bool) else ("Yes" if val else "No"))

    # Security Recommendations
    document.add_heading('Security Recommendations', level=1)
    if prediction.lower() == 'safe':
        rec = "The scan suggests this content is low risk. However, always exercise caution. Do not share passwords or OTPs even on safe-looking sites."
    elif prediction.lower() == 'suspicious':
        rec = "TAKE CAUTION: This content displays patterns commonly used in phishing campaigns. Do NOT interact with links or download attachments from this source."
    else:
        rec = "IMMEDIATE ACTION REQUIRED: This content has been flagged as high-risk PHISHING. If this was an email, delete it immediately. If it was a website, close it and clear your browser cache."
    
    document.add_paragraph(rec)

    # Reporting
    document.add_heading('How to Report', level=2)
    document.add_paragraph("Official Portal: https://cybercrime.gov.in (India हेल्पलाइन: 1930)")
    document.add_paragraph("Google Report Phish: https://safebrowsing.google.com/safebrowsing/report_phish/")

    # Footer
    document.add_paragraph()
    footer_p = document.add_paragraph("Report generated by PhishGuard AI Core | © 2025 Kunal Navdhinge")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    target = io.BytesIO()
    document.save(target)
    target.seek(0)
    return target


# ──────────────────────────────────────────────────────────────
# API Routes
# ──────────────────────────────────────────────────────────────

@app.route('/')
def serve_frontend():
    """Serve the main frontend page."""
    return send_from_directory(app.static_folder, 'index.html')


@app.route('/<path:path>')
def serve_static(path):
    """Serve static frontend files."""
    return send_from_directory(app.static_folder, path)


@app.route('/api/scan-url', methods=['POST'])
@limiter.limit("30 per minute")
def scan_url():
    """
    Scan a URL for phishing detection.
    Expects JSON: {"url": "https://example.com"}
    """
    try:
        data = request.get_json()
        if not data or 'url' not in data:
            return jsonify({'error': 'URL is required'}), 400

        url = data['url'].strip()
        if not is_valid_url(url):
            return jsonify({'error': 'Invalid URL format'}), 400

        url_clean = sanitize_input(url)
        if not url_clean:
            return jsonify({'error': 'Invalid URL after sanitization'}), 400

        result = perform_scan(url_clean, source='URL Input')
        return jsonify(result), 200

    except RuntimeError as e:
        return jsonify({'error': str(e)}), 503
    except Exception as e:
        print(f"[ERROR] Scan failed: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/scan-qr', methods=['POST'])
@limiter.limit("20 per minute")
def scan_qr():
    """
    Scan a URL extracted from a QR code.
    The QR decoding happens on the frontend; this receives the decoded URL.
    Expects JSON: {"url": "https://example.com", "source": "qr_camera" | "qr_image"}
    """
    try:
        data = request.get_json()
        if not data or 'url' not in data:
            return jsonify({'error': 'URL from QR code is required'}), 400

        url = data['url'].strip()
        if not is_valid_url(url):
            return jsonify({'error': 'Invalid URL extracted from QR code'}), 400

        url_clean = sanitize_input(url)
        if not url_clean:
            return jsonify({'error': 'Invalid URL after sanitization'}), 400

        source = data.get('source', 'QR Code')
        source_label = 'QR Code (Camera)' if source == 'qr_camera' else 'QR Code (Image)'

        result = perform_scan(url_clean, source=source_label)
        return jsonify(result), 200

    except RuntimeError as e:
        return jsonify({'error': str(e)}), 503
    except Exception as e:
        print(f"[ERROR] QR Scan failed: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/scan-email', methods=['POST'])
@limiter.limit("20 per minute")
def scan_email():
    """
    Scan an email body for phishing indicators.
    Expects JSON: {"content": "..."}
    """
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'Email content is required'}), 400

        content = data['content'].strip()
        if len(content) < 10:
            return jsonify({'error': 'Email content is too short'}), 400
        if len(content) > 50000:
            return jsonify({'error': 'Email content is too long'}), 400

        result = analyze_email_text(content)
        return jsonify(result), 200

    except Exception as e:
        print(f"[ERROR] Email Scan failed: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/generate-report', methods=['POST'])
@limiter.limit("15 per minute")
def generate_report():
    """
    Generate a DOCX report from scan data.
    Expects JSON with full scan result data.
    """
    try:
        data = request.get_json()
        if not data or 'url' not in data:
            return jsonify({'error': 'Scan data is required'}), 400

        docx_buffer = generate_docx_report(data)

        # Create safe filename
        safe_url = re.sub(r'[^a-zA-Z0-9]', '_', data.get('url', 'scan'))[:40]
        filename = f'PhishGuard_Report_{safe_url}.docx'

        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"[ERROR] DOCX generation failed: {e}")
        return jsonify({'error': 'Failed to generate DOCX report'}), 500


@app.route('/api/report-url', methods=['POST'])
@limiter.limit("10 per minute")
def report_url():
    """
    Report a phishing URL.
    Expects JSON: {"url": "...", "reason": "...", "report_to_cybercrime": true/false}
    """
    try:
        data = request.get_json()
        if not data or 'url' not in data:
            return jsonify({'error': 'URL is required'}), 400

        url = sanitize_input(data.get('url', ''))
        reason = sanitize_input(data.get('reason', 'No reason provided'))
        report_to_cybercrime = data.get('report_to_cybercrime', False)

        if not is_valid_url(url):
            return jsonify({'error': 'Invalid URL format'}), 400

        report_data = {
            'url': url,
            'reason': reason,
            'reported_at': datetime.utcnow().isoformat() + 'Z',
            'status': 'pending',
            'report_to_cybercrime': report_to_cybercrime
        }

        doc_id = save_report(report_data)
        if doc_id:
            report_data['report_id'] = doc_id
        else:
            in_memory_reports.insert(0, report_data)
            if len(in_memory_reports) > 100:
                in_memory_reports.pop()

        # Cybercrime reporting portals info
        cybercrime_links = {
            'india_portal': 'https://cybercrime.gov.in',
            'india_helpline': '1930',
            'google_safe_browsing': 'https://safebrowsing.google.com/safebrowsing/report_phish/',
            'phishtank': 'https://www.phishtank.com/',
            'apwg_email': 'reportphishing@apwg.org'
        }

        return jsonify({
            'message': 'URL reported successfully',
            'report': report_data,
            'cybercrime_links': cybercrime_links
        }), 201

    except Exception as e:
        print(f"[ERROR] Report failed: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/recent-scans', methods=['GET'])
@limiter.limit("60 per minute")
def recent_scans():
    """Get recent scan results."""
    try:
        limit = min(int(request.args.get('limit', 20)), 50)
        scans = get_recent_scans(limit)
        if not scans:
            scans = in_memory_scans[:limit]
        return jsonify({'scans': scans}), 200
    except Exception as e:
        print(f"[ERROR] Failed to fetch recent scans: {e}")
        return jsonify({'scans': in_memory_scans[:20]}), 200


@app.route('/api/stats', methods=['GET'])
@limiter.limit("60 per minute")
def stats():
    """Get scan statistics."""
    try:
        statistics = get_scan_stats()
        return jsonify(statistics), 200
    except Exception as e:
        print(f"[ERROR] Failed to fetch stats: {e}")
        return jsonify(in_memory_stats), 200


@app.route('/api/email-samples', methods=['GET'])
def get_email_samples():
    """Serve email samples for testing."""
    try:
        samples_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'samples', 'email_samples.json')
        if os.path.exists(samples_path):
            with open(samples_path, 'r', encoding='utf-8') as f:
                samples = json.load(f)
            return jsonify(samples), 200
        else:
            return jsonify({'error': 'Samples file not found'}), 404
    except Exception as e:
        print(f"[ERROR] Failed to fetch email samples: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'model_loaded': model is not None,
        'timestamp': datetime.utcnow().isoformat() + 'Z'
    }), 200


# ──────────────────────────────────────────────────────────────
# Error Handlers
# ──────────────────────────────────────────────────────────────

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found'}), 404


@app.errorhandler(429)
def rate_limited(e):
    return jsonify({'error': 'Rate limit exceeded. Please try again later.'}), 429


@app.errorhandler(500)
def server_error(e):
    return jsonify({'error': 'Internal server error'}), 500


# ──────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────
if __name__ == '__main__':
    init_db()
    load_model()

    port = int(os.getenv('PORT', 5000))
    debug = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'

    print(f"\n{'='*60}")
    print(f"  PhishGuard — Phishing Detection API Server")
    print(f"  Running on http://localhost:{port}")
    print(f"  Model loaded: {model is not None}")
    print(f"  Features: QR Scan, DOCX Reports, Cybercrime Reporting")
    print(f"{'='*60}\n")

    app.run(host='0.0.0.0', port=port, debug=debug)
